import io
import re
from collections import Counter

from pypdf import PdfReader
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer, ENGLISH_STOP_WORDS
from sklearn.metrics.pairwise import cosine_similarity


def extract_contact_info(text):
    """Extract email addresses and phone numbers from resume text.

    Returns:
        Dict with 'email' (str or '') and 'phone' (str or '').
    """
    # Email: standard pattern
    email_match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    email = email_match.group(0) if email_match else ""

    # Phone: flexible pattern for various formats
    # Matches: +1-234-567-8901, (234) 567-8901, 234-567-8901, 234.567.8901, 2345678901
    phone_match = re.search(
        r"(?:\+?\d{1,3}[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}", text
    )
    phone = phone_match.group(0).strip() if phone_match else ""

    return {"email": email, "phone": phone}


def extract_experience_years(text):
    """Extract years of experience mentioned in resume text.

    Looks for patterns like '5 years', '3+ years of experience',
    '5-7 years experience', etc.

    Returns:
        Float representing the max years found, or 0.0 if none detected.
    """
    patterns = [
        r"(\d{1,2})\+?\s*(?:years|yrs)[\s\w]*(?:experience|exp)",
        r"(\d{1,2})\s*-\s*\d{1,2}\s*(?:years|yrs)",
        r"(\d{1,2})\+?\s*(?:years|yrs)",
    ]
    years_found = []
    text_lower = text.lower()
    for pattern in patterns:
        matches = re.findall(pattern, text_lower)
        years_found.extend(int(m) for m in matches)

    return float(max(years_found)) if years_found else 0.0


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF file bytes."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception:
        return ""


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file bytes, including table content."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception:
        return ""


def extract_text(uploaded_file):
    """Route to correct extractor based on file extension.

    Returns:
        Tuple of (filename, extracted_text).
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        text = ""

    return uploaded_file.name, text


def calculate_tfidf_similarity(job_description, resume_texts):
    """Calculate TF-IDF cosine similarity between a job description and resumes.

    Args:
        job_description: Job description string.
        resume_texts: List of resume text strings.

    Returns:
        List of float scores (0.0 to 1.0), one per resume.
    """
    corpus = [job_description] + resume_texts

    vectorizer = TfidfVectorizer(
        stop_words="english",
        max_features=5000,
        ngram_range=(1, 2),
        min_df=1,
    )
    tfidf_matrix = vectorizer.fit_transform(corpus)

    similarities = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:])
    return similarities.flatten().tolist()


def extract_keywords(text, top_n=30):
    """Extract top keywords from text using frequency analysis.

    Returns:
        Set of top keyword strings.
    """
    text_clean = re.sub(r"[^a-zA-Z\s]", " ", text.lower())
    words = text_clean.split()
    words = [w for w in words if len(w) > 2]

    filler = {
        "also", "well", "using", "used", "work", "working", "able",
        "good", "new", "need", "required", "including", "etc",
        "experience", "skills", "years", "strong", "knowledge",
    }
    freq = Counter(words)
    keywords = {
        w: c for w, c in freq.items()
        if w not in ENGLISH_STOP_WORDS and w not in filler
    }

    sorted_kw = sorted(keywords.items(), key=lambda x: x[1], reverse=True)
    return set(w for w, _ in sorted_kw[:top_n])


def keyword_overlap_analysis(job_description, resume_text, custom_jd_keywords=None):
    """Compute keyword overlap between a job description and a resume.

    Args:
        job_description: Job description text (used if custom_jd_keywords is None).
        resume_text: Resume text string.
        custom_jd_keywords: Optional set of user-edited keywords to use instead
            of auto-extracting from job_description.

    Returns:
        Dict with matched_keywords, missing_keywords, and overlap_score.
    """
    jd_keywords = custom_jd_keywords if custom_jd_keywords is not None else extract_keywords(job_description)
    resume_keywords = extract_keywords(resume_text)

    matched = jd_keywords & resume_keywords
    missing = jd_keywords - resume_keywords

    overlap_score = len(matched) / len(jd_keywords) if jd_keywords else 0.0

    return {
        "jd_keywords": jd_keywords,
        "resume_keywords": resume_keywords,
        "matched_keywords": matched,
        "missing_keywords": missing,
        "overlap_score": overlap_score,
    }
